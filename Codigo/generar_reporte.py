# Codigo/generar_reporte.py

import os
import sys

# Agregar el directorio raíz al path para poder importar Codigo
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import torch
from torchvision import transforms
from PIL import Image, ImageDraw, ImageFont
from collections import Counter
from docx import Document
from docx.shared import Inches
from Codigo import config
from Codigo.modelo_cnn import CNNBasica
import matplotlib.pyplot as plt
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader

# Cargar modelo entrenado
def cargar_modelo(ruta_modelo, num_clases):
    modelo = CNNBasica(num_classes=num_clases)
    modelo.load_state_dict(torch.load(ruta_modelo, map_location=torch.device('cpu')))
    modelo.eval()
    return modelo

# Calcular avance de obra
def calcular_avance(conteo):
    max_orden = 0
    etapa_actual = "Desconocida"
    avance_porcentaje = 0

    if not config.ETAPAS:
        return "No configurado", 0

    clases_detectadas = conteo.keys()

    for clase in clases_detectadas:
        # Busca coincidencia parcial o exacta en las claves de config.ETAPAS
        # Por ejemplo, si la clase es 'Zapata_Frente', busca si contiene 'Zapata'
        for key_etapa, info in config.ETAPAS.items():
            if key_etapa.lower() in clase.lower():
                if info["orden"] > max_orden:
                    max_orden = info["orden"]
                    etapa_actual = key_etapa
                    avance_porcentaje = info["avance"]

    return etapa_actual, avance_porcentaje

# Clasificar imágenes de "Datos/"
def clasificar_imagenes(modelo, transform, clases):
    resultados = []
    for nombre_archivo in os.listdir(config.DATOS_DIR):
        if nombre_archivo.lower().endswith(('.jpg', '.jpeg', '.png')):
            ruta = os.path.join(config.DATOS_DIR, nombre_archivo)
            try:
                img = Image.open(ruta).convert('RGB')
                entrada = transform(img).unsqueeze(0)
                salida = modelo(entrada)
                _, pred = torch.max(salida, 1)
                idx = pred.item()
                if idx < len(clases):
                    etiqueta = clases[idx]
                    resultados.append((nombre_archivo, etiqueta, ruta))
                else:
                    print(f"⚠️ Advertencia: El modelo predijo la clase {idx}, pero solo hay {len(clases)} clases definidas.")
            except Exception as e:
                print(f"⚠️ Error con {nombre_archivo}: {e}")
    return resultados

# Graficar conteo de clases
def guardar_grafica_conteo(conteo, ruta):
    clases = list(conteo.keys())
    cantidades = list(conteo.values())

    plt.figure(figsize=(8, 5))
    plt.bar(clases, cantidades, color='skyblue')
    plt.xlabel('Clase')
    plt.ylabel('Cantidad')
    plt.title('Distribución de predicciones')
    plt.tight_layout()
    plt.savefig(ruta)
    plt.close()

# Crear documento Word
def generar_word(resultados, conteo):
    doc = Document()
    doc.add_heading('Reporte de Clasificación y Avance de Obra', level=1)

    # Calcular avance
    etapa, porcentaje = calcular_avance(conteo)

    doc.add_heading('Resumen de Avance', level=2)
    p = doc.add_paragraph()
    p.add_run(f'Etapa actual detectada: ').bold = True
    p.add_run(f'{etapa}\n')
    p.add_run(f'Avance estimado de obra: ').bold = True
    p.add_run(f'{porcentaje}%')

    total = sum(conteo.values())
    doc.add_heading('Detalles de Clasificación', level=2)
    doc.add_paragraph(f'Total de imágenes clasificadas: {total}')

    tabla = doc.add_table(rows=1, cols=2)
    tabla.style = 'Table Grid'
    tabla.cell(0, 0).text = 'Clase'
    tabla.cell(0, 1).text = 'Cantidad'

    for clase, cant in conteo.items():
        row = tabla.add_row()
        row.cells[0].text = clase
        row.cells[1].text = str(cant)

    doc.add_picture(os.path.join(config.RESULTADOS_DIR, "grafica_conteo.png"), width=Inches(5))
    doc.add_page_break()

    for nombre, etiqueta, ruta in resultados:
        doc.add_paragraph(f'{nombre} → Clase: {etiqueta}')
        img = Image.open(ruta)
        marcada_path = os.path.join(config.RESULTADOS_DIR, f"marcada_{nombre}")
        draw = ImageDraw.Draw(img)
        draw.text((10, 10), etiqueta, fill=(255, 0, 0), font=ImageFont.load_default())
        img.save(marcada_path)
        doc.add_picture(marcada_path, width=Inches(3.5))

    output_path = os.path.join(config.RESULTADOS_DIR, 'reporte_resultados.docx')
    doc.save(output_path)
    print(f"📝 Word guardado: {output_path}")

# Crear PDF con resultados
def generar_pdf(resultados, conteo):
    pdf_path = os.path.join(config.RESULTADOS_DIR, "reporte_resultados.pdf")
    c = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter

    # Calcular avance
    etapa, porcentaje = calcular_avance(conteo)

    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, height - 50, "Reporte de Avance de Obra")

    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, height - 80, f"Etapa actual: {etapa}")
    c.drawString(300, height - 80, f"Avance estimado: {porcentaje}%")

    c.setFont("Helvetica", 12)
    c.drawString(50, height - 110, f"Total imágenes procesadas: {sum(conteo.values())}")

    y = height - 140
    for clase, cant in conteo.items():
        c.drawString(60, y, f"{clase}: {cant}")
        y -= 20

    c.drawImage(os.path.join(config.RESULTADOS_DIR, "grafica_conteo.png"), 50, y - 250, width=500, preserveAspectRatio=True, mask='auto')
    c.showPage()

    for nombre, etiqueta, ruta in resultados:
        c.setFont("Helvetica", 12)
        c.drawString(50, height - 50, f"{nombre} → {etiqueta}")
        marcada_path = os.path.join(config.RESULTADOS_DIR, f"marcada_{nombre}")
        if os.path.exists(marcada_path):
            c.drawImage(ImageReader(marcada_path), 50, height - 350, width=300, preserveAspectRatio=True, mask='auto')
        c.showPage()

    c.save()
    print(f"📄 PDF guardado: {pdf_path}")

# MAIN
if __name__ == "__main__":
    os.makedirs(config.RESULTADOS_DIR, exist_ok=True)

    # Verificar si hay imágenes en Datos
    if not os.path.exists(config.DATOS_DIR) or not any(f.lower().endswith(('.jpg', '.jpeg', '.png')) for f in os.listdir(config.DATOS_DIR)):
        print(f"⚠️ Alerta: No se encontraron imágenes en {config.DATOS_DIR}. Asegúrate de poner las imágenes a clasificar en la carpeta 'Datos'.")
        sys.exit()

    transform = transforms.Compose([
        transforms.Resize(config.IMAGE_SIZE),
        transforms.ToTensor(),
        transforms.Normalize([0.5]*3, [0.5]*3)
    ])

    modelo = cargar_modelo(os.path.join(config.MODELOS_DIR, 'modelo_cnn.pth'), config.NUM_CLASSES)
    # nombres de carpetas = clases
    clases = sorted([d for d in os.listdir(config.IMAGENES_DIR) if os.path.isdir(os.path.join(config.IMAGENES_DIR, d))])

    resultados = clasificar_imagenes(modelo, transform, clases)
    conteo = Counter([etq for _, etq, _ in resultados])

    guardar_grafica_conteo(conteo, os.path.join(config.RESULTADOS_DIR, "grafica_conteo.png"))
    generar_word(resultados, conteo)
    generar_pdf(resultados, conteo)
